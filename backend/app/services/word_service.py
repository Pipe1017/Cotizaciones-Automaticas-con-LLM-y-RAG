"""
Genera los dos documentos Word por cotización:
  - carta.docx  — carta formal de presentación
  - cotizacion.docx — tabla de ítems, totales y condiciones comerciales
"""
import io
import os
from datetime import date
from decimal import Decimal

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ──────────────────────────────────────────────
NAVY   = "0F2560"
ORANGE = "F97316"
WHITE  = "FFFFFF"
LIGHT  = "F1F5F9"

# ── Asesor directory ───────────────────────────────────────────
ASESORES = {
    "Aura María Gallego": {
        "cargo": "Dirección Comercial",
        "email": "aura@opex.com.co",
        "telefono": "(57) 311 312 8684",
    },
    "Juan David Giraldo": {
        "cargo": "Asesor Comercial",
        "email": "juan.giraldo@opex.com.co",
        "telefono": "",
    },
    "Alejandro Rendón": {
        "cargo": "Asesor Comercial",
        "email": "alejandro.rendon@opex.com.co",
        "telefono": "",
    },
    "Diego Arboleda": {
        "cargo": "Asesor Comercial",
        "email": "diego.arboleda@opex.com.co",
        "telefono": "",
    },
}

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "static", "Opex.png")

CIUDADES = {
    "med": "Medellín", "bog": "Bogotá", "cal": "Cali",
    "bar": "Barranquilla", "buc": "Bucaramanga", "man": "Manizales",
    "per": "Pereira", "car": "Cartagena", "cuc": "Cúcuta",
}


# ── XML helpers ────────────────────────────────────────────────
def _set_cell_color(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), kwargs.get(side, "none"))
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "D1D5DB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _para(cell, text, bold=False, color=None, size=10, align=None):
    p = cell.paragraphs[0]
    p.clear()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_header(doc: Document, numero: str, fecha, ciudad_code: str):
    """Header con logo + número de cotización."""
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(9)
    table.columns[2].width = Cm(4)

    # Logo
    logo_cell = table.cell(0, 0)
    _set_cell_color(logo_cell, WHITE)
    if os.path.exists(LOGO_PATH):
        logo_cell.paragraphs[0].add_run().add_picture(LOGO_PATH, width=Cm(3.2))

    # Center: company tagline
    center = table.cell(0, 1)
    _set_cell_color(center, NAVY)
    cp = center.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("www.opex.com.co")
    cr.font.color.rgb = RGBColor.from_string(WHITE)
    cr.font.size = Pt(9)

    # Right: N° + fecha + ciudad
    right = table.cell(0, 2)
    _set_cell_color(right, NAVY)
    ciudad = CIUDADES.get(ciudad_code, ciudad_code.title())
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(f"N° {numero}\n{fecha}\n{ciudad}")
    rr.font.color.rgb = RGBColor.from_string(WHITE)
    rr.font.size = Pt(8)
    rr.bold = True

    doc.add_paragraph()


def _add_footer_bar(doc: Document):
    """Barra inferior naranja + lema."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(14)
    _set_cell_color(table.cell(0, 0), NAVY)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("H₂")
    r.font.color.rgb = RGBColor.from_string(WHITE)
    r.font.bold = True
    right = table.cell(0, 1)
    _set_cell_color(right, ORANGE)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Innovación, energía y movimiento")
    r2.font.color.rgb = RGBColor.from_string(WHITE)
    r2.font.bold = True
    r2.font.size = Pt(10)


def _asesor_info(asesor_nombre: str) -> dict:
    return ASESORES.get(asesor_nombre, {
        "cargo": "Asesor Comercial",
        "email": "info@opex.com.co",
        "telefono": "",
    })


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 1 — Carta formal
# ══════════════════════════════════════════════════════════════
def generate_carta(data: dict) -> bytes:
    """
    data keys: cliente, contacto_nombre, numero_cotizacion, fecha,
               ciudad_cotizacion, titulo_oportunidad, asesor, business_line_nombre
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    numero   = data.get("numero_cotizacion", "")
    fecha    = data.get("fecha", date.today())
    if isinstance(fecha, date):
        fecha_str = fecha.strftime("%d de %B de %Y")
    else:
        fecha_str = str(fecha)

    ciudad_code = data.get("ciudad_cotizacion", "med")
    cliente     = data.get("cliente", "Cliente")
    contacto    = data.get("contacto_nombre") or ""
    asesor_nom  = data.get("asesor") or "Aura María Gallego"
    asesor      = _asesor_info(asesor_nom)
    bl_nombre   = data.get("business_line_nombre", "")
    titulo_op   = data.get("titulo_oportunidad", f"Suministro {bl_nombre}")

    # ── Header ──────────────────────────────────────────────
    _add_header(doc, numero, fecha_str, ciudad_code)

    # ── Fecha y destinatario ─────────────────────────────────
    doc.add_paragraph(fecha_str)

    p_dest = doc.add_paragraph()
    p_dest.add_run(f"{contacto}\n" if contacto else "").bold = False
    p_dest.add_run(cliente).bold = True
    p_dest.add_run(f"\n{CIUDADES.get(ciudad_code, ciudad_code.title())}")

    doc.add_paragraph()

    # Asunto
    p_asunto = doc.add_paragraph()
    r = p_asunto.add_run("Asunto: ")
    r.bold = True
    p_asunto.add_run(f"Presentación de oferta — {titulo_op}")

    doc.add_paragraph()

    # Saludo
    nombre_corto = contacto.split()[0] if contacto else "equipo"
    doc.add_paragraph(f"Estimado/a {nombre_corto},")
    doc.add_paragraph()

    # Cuerpo
    body = (
        f"A continuación, nos permitimos presentar la oferta económica correspondiente al "
        f"requerimiento de {bl_nombre.lower() if bl_nombre else 'productos y servicios'} para "
        f"{cliente}, con la confianza de ofrecer soluciones de alta calidad, respaldo técnico "
        f"y condiciones competitivas adaptadas a sus necesidades operativas.\n\n"
        "Quedamos completamente atentos a resolver cualquier duda, inquietud o requerimiento "
        "adicional que pueda surgir durante su proceso de evaluación. Es para nosotros un "
        "honor poder acompañarlos en la optimización de sus operaciones y esperamos que "
        "nuestra propuesta sea de su entera satisfacción.\n\n"
        "Agradecemos la oportunidad de participar en este proceso y quedamos a su disposición."
    )
    doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_paragraph("Cordialmente,")
    doc.add_paragraph()
    doc.add_paragraph()

    # Firma
    p_firma = doc.add_paragraph()
    r_nom = p_firma.add_run(asesor_nom)
    r_nom.bold = True
    r_nom.font.size = Pt(11)
    r_nom.font.color.rgb = RGBColor.from_string(NAVY)

    p_cargo = doc.add_paragraph(asesor["cargo"])
    p_cargo.runs[0].font.color.rgb = RGBColor.from_string(NAVY)

    contacto_line = f"www.opex.com.co  |  {asesor['email']}"
    if asesor["telefono"]:
        contacto_line += f"  |  {asesor['telefono']}"
    doc.add_paragraph(contacto_line)

    doc.add_paragraph()
    _add_footer_bar(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 2 — Cotización (ítems + totales + condiciones)
# ══════════════════════════════════════════════════════════════
def generate_cotizacion(data: dict) -> bytes:
    """
    data keys: cliente, contacto_nombre, numero_cotizacion, fecha,
               ciudad_cotizacion, asesor, items[], subtotal_usd,
               iva_pct, total_usd, observaciones, condiciones_*,
               validez_oferta, fecha_entrega
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(1.5)

    numero      = data.get("numero_cotizacion", "")
    fecha       = data.get("fecha", date.today())
    if isinstance(fecha, date):
        fecha_str = fecha.strftime("%d/%m/%Y")
    else:
        fecha_str = str(fecha)

    ciudad_code = data.get("ciudad_cotizacion", "med")
    cliente     = data.get("cliente", "Cliente")
    contacto    = data.get("contacto_nombre") or "—"
    asesor_nom  = data.get("asesor") or "Aura María Gallego"
    asesor      = _asesor_info(asesor_nom)
    items       = data.get("items", [])
    subtotal    = float(data.get("subtotal_usd") or 0)
    iva_pct     = float(data.get("iva_pct") or 19)
    total       = float(data.get("total_usd") or 0)

    # ── Header ──────────────────────────────────────────────
    _add_header(doc, numero, fecha_str, ciudad_code)

    # ── Info cliente ─────────────────────────────────────────
    info = doc.add_table(rows=2, cols=4)
    info.autofit = False
    col_w = [Cm(3), Cm(6), Cm(3), Cm(5)]
    for i, w in enumerate(col_w):
        info.columns[i].width = w

    labels = [("CLIENTE", cliente), ("CONTACTO", contacto),
              ("N° COTIZACIÓN", numero), ("FECHA", fecha_str)]
    for idx, (lbl, val) in enumerate(labels):
        row = idx // 2
        col = (idx % 2) * 2
        lc = info.cell(row, col)
        vc = info.cell(row, col + 1)
        _set_cell_color(lc, NAVY)
        _para(lc, lbl, bold=True, color=WHITE, size=8)
        _para(vc, val, size=9)

    doc.add_paragraph()

    # ── Tabla de ítems ───────────────────────────────────────
    COLS = ["#", "Referencia", "Descripción", "Cód. SAP", "Marca", "Cant.", "P. Unit. USD", "P. Total USD"]
    WIDTHS = [Cm(0.7), Cm(2.2), Cm(6.5), Cm(1.8), Cm(1.8), Cm(1.0), Cm(2.2), Cm(2.2)]

    normales   = [i for i in items if not i.get("opcional", False)]
    opcionales = [i for i in items if i.get("opcional", False)]

    tbl = doc.add_table(rows=1 + len(normales), cols=8)
    tbl.autofit = False
    for i, w in enumerate(WIDTHS):
        tbl.columns[i].width = w

    # Header row
    hdr = tbl.rows[0]
    for i, (col_name, cell) in enumerate(zip(COLS, hdr.cells)):
        _set_cell_color(cell, NAVY)
        _para(cell, col_name, bold=True, color=WHITE, size=8,
              align=WD_ALIGN_PARAGRAPH.CENTER if i >= 5 else WD_ALIGN_PARAGRAPH.LEFT)

    # Data rows — solo normales
    for r_idx, item in enumerate(normales, start=1):
        row = tbl.rows[r_idx]
        bg = WHITE if r_idx % 2 == 0 else LIGHT
        vals = [
            str(item.get("item_number", r_idx)),
            item.get("referencia_usa") or "—",
            item.get("descripcion", ""),
            item.get("referencia_cod_proveedor") or "—",
            item.get("marca") or "—",
            str(item.get("cantidad", 1)),
            f"${float(item.get('precio_unitario_usd', 0)):,.2f}",
            f"${float(item.get('precio_total_usd', 0)):,.2f}",
        ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT,   WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.RIGHT]
        for ci, (val, aln, cell) in enumerate(zip(vals, aligns, row.cells)):
            _set_cell_color(cell, bg)
            _para(cell, val, size=8, align=aln)

    # Tabla de OPCIONALES
    if opcionales:
        doc.add_paragraph()
        ph = doc.add_paragraph("ÍTEMS OPCIONALES — No incluidos en el total")
        ph.runs[0].bold = True
        ph.runs[0].font.size = Pt(9)
        ph.runs[0].font.color.rgb = RGBColor(0x0F, 0x25, 0x60)

        otbl = doc.add_table(rows=1 + len(opcionales), cols=8)
        otbl.autofit = False
        for i, w in enumerate(WIDTHS):
            otbl.columns[i].width = w

        ohdr = otbl.rows[0]
        OPT_COLS = ["#", "REFERENCIA", "DESCRIPCIÓN", "COD. PROVEEDOR", "MARCA", "OPCIONAL", "PRECIO REF.", "—"]
        for i, (col_name, cell) in enumerate(zip(OPT_COLS, ohdr.cells)):
            _set_cell_color(cell, "64748B")
            _para(cell, col_name, bold=True, color=WHITE, size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i >= 5 else WD_ALIGN_PARAGRAPH.LEFT)

        for r_idx, item in enumerate(opcionales, start=1):
            row = otbl.rows[r_idx]
            nota = item.get("notas") or ""
            desc = item.get("descripcion", "")
            desc_full = f"{desc}\n[{nota}]" if nota else desc
            vals = [
                f"OPC{r_idx}",
                item.get("referencia_usa") or "—",
                desc_full,
                item.get("referencia_cod_proveedor") or "—",
                item.get("marca") or "—",
                "OPCIONAL",
                f"${float(item.get('precio_unitario_usd', 0)):,.2f}",
                "—",
            ]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.LEFT,   WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.CENTER]
            for ci, (val, aln, cell) in enumerate(zip(vals, aligns, row.cells)):
                _set_cell_color(cell, "F1F5F9")
                _para(cell, val, size=8, color="64748B", align=aln)

    doc.add_paragraph()

    # ── Servicios de Ingeniería ───────────────────────────────
    servicios = data.get("servicios", [])
    if servicios:
        ph = doc.add_paragraph("SERVICIOS DE INGENIERÍA")
        ph.runs[0].bold = True
        ph.runs[0].font.size = Pt(9)
        ph.runs[0].font.color.rgb = RGBColor(0x0F, 0x25, 0x60)

        stbl = doc.add_table(rows=1 + len(servicios) + 1, cols=4)
        stbl.autofit = False
        for i, w in enumerate([Cm(4.5), Cm(8), Cm(1.5), Cm(3.5)]):
            stbl.columns[i].width = w

        for i, txt in enumerate(["ROL", "MOTIVO / DESCRIPCIÓN", "HORAS", "SUBTOTAL"]):
            _set_cell_color(stbl.rows[0].cells[i], "64748B")
            _para(stbl.rows[0].cells[i], txt, bold=True, color=WHITE, size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT)

        svc_sub = Decimal("0")
        for r_idx, svc in enumerate(servicios, start=1):
            row = stbl.rows[r_idx]
            sub = float(svc.get("subtotal_usd", 0))
            svc_sub += Decimal(str(sub))
            bg = WHITE if r_idx % 2 == 0 else LIGHT
            for c in row.cells:
                _set_cell_color(c, bg)
            _para(row.cells[0], svc.get("nombre", ""), bold=True, size=8)
            _para(row.cells[1], svc.get("motivo", ""), size=8, color="64748B")
            _para(row.cells[2], str(svc.get("horas", 0)), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            _para(row.cells[3], f"${sub:,.2f}", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

        tot_row = stbl.rows[-1]
        _set_cell_color(tot_row.cells[0], LIGHT)
        _set_cell_color(tot_row.cells[1], LIGHT)
        _set_cell_color(tot_row.cells[2], LIGHT)
        _set_cell_color(tot_row.cells[3], LIGHT)
        _para(tot_row.cells[2], "Subtotal", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _para(tot_row.cells[3], f"${float(svc_sub):,.2f}", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

        doc.add_paragraph()

    # ── Totales ──────────────────────────────────────────────
    svc_subtotal = float(data.get("servicios_subtotal_usd", 0))
    totals_rows = [("Subtotal Productos USD", f"${subtotal:,.2f}")]
    if svc_subtotal > 0:
        totals_rows.append(("Subtotal Servicios USD", f"${svc_subtotal:,.2f}"))
    totals_rows += [
        (f"IVA {iva_pct:.0f}%", f"${(subtotal + svc_subtotal) * iva_pct / 100:,.2f}"),
        ("TOTAL USD",           f"${total:,.2f}"),
    ]

    ttbl = doc.add_table(rows=len(totals_rows), cols=2)
    ttbl.autofit = False
    ttbl.columns[0].width = Cm(14)
    ttbl.columns[1].width = Cm(3.5)

    totals = totals_rows
    for i, (lbl, val) in enumerate(totals):
        lc, vc = ttbl.rows[i].cells
        is_total = (lbl == "TOTAL USD")
        bg = NAVY if is_total else LIGHT
        tc = WHITE if is_total else "374151"
        _set_cell_color(lc, bg)
        _set_cell_color(vc, bg)
        _para(lc, lbl, bold=is_total, color=tc, size=9,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
        _para(vc, val, bold=True, color=tc, size=9,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()

    # ── Condiciones comerciales ───────────────────────────────
    conds = [
        ("Validez de la oferta", data.get("validez_oferta") or "30 días"),
        ("Condiciones de pago",  data.get("condiciones_pago") or "Las habituales"),
        ("Entrega estimada",     data.get("fecha_entrega") or data.get("condiciones_entrega") or "—"),
        ("Garantía",             data.get("condiciones_garantia") or "—"),
    ]

    p_cond = doc.add_paragraph()
    r = p_cond.add_run("CONDICIONES COMERCIALES")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.size = Pt(9)

    ctbl = doc.add_table(rows=len(conds), cols=2)
    ctbl.autofit = False
    ctbl.columns[0].width = Cm(5)
    ctbl.columns[1].width = Cm(12.5)
    for i, (lbl, val) in enumerate(conds):
        lc, vc = ctbl.rows[i].cells
        _set_cell_color(lc, NAVY)
        _set_cell_color(vc, LIGHT if i % 2 == 0 else WHITE)
        _para(lc, lbl, bold=True, color=WHITE, size=8)
        _para(vc, val, size=8)

    # Observaciones
    obs = data.get("observaciones")
    if obs:
        doc.add_paragraph()
        p_obs = doc.add_paragraph()
        r2 = p_obs.add_run("Observaciones: ")
        r2.bold = True
        r2.font.size = Pt(8)
        p_obs.add_run(obs).font.size = Pt(8)

    doc.add_paragraph()

    # Pie con contacto asesor
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie_text = f"{asesor_nom}  |  {asesor['email']}"
    if asesor["telefono"]:
        pie_text += f"  |  {asesor['telefono']}"
    pie_text += "  |  www.opex.com.co"
    r_pie = p_pie.add_run(pie_text)
    r_pie.font.size = Pt(8)
    r_pie.font.color.rgb = RGBColor.from_string(NAVY)

    _add_footer_bar(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
